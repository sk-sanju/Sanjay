from docx import Document

# Create a new Document
doc = Document()

# Title and Contact Information
doc.add_heading('Sanjay S', 0)
doc.add_paragraph(
    'Phone: +91 9539251789 | Email: sanjayskpy1@gmail.com\n'
    'LinkedIn: www.linkedin.com/in/sanjay-s953925 | Portfolio: https://sk-sanju.github.io/Sanjay/'
)

# Career Objective
doc.add_heading('Career Objective', level=1)
doc.add_paragraph(
    "As an aspiring Data Analyst, I am passionate about utilizing my technical expertise in SQL, Python, Power BI, and Tableau "
    "to analyze and visualize data, driving insights and data-driven decision-making. I aim to apply my skills to contribute effectively "
    "to organizational growth and problem-solving."
)

# Core Skills
doc.add_heading('Core Skills', level=1)
doc.add_paragraph(
    "- SQL: Expertise in querying and managing relational databases.\n"
    "- Oracle: Familiar with writing SQL queries, performing data retrieval, and working with Oracle SQL Developer for relational database tasks.\n"
    "- Power BI: Skilled in creating interactive dashboards and reports to visualize key metrics.\n"
    "- Advanced Excel: Strong command of data analysis, pivot tables, and complex formulas.\n"
    "- Python: Proficient in data analysis and automation using Python libraries like Pandas, NumPy, and Matplotlib."
)

# Education
doc.add_heading('Education', level=1)
doc.add_paragraph(
    "Bachelor of Computer Application (BCA)\n"
    "Sadanam Kumaran College, Pathiripala, Palakkad / University of Calicut — 2021 – 2023\n"
    "Key Coursework: Python, Database Management Systems (DBMS), Structured Query Language (SQL), Data Structures & Algorithms (DSA)"
)

# Certifications
doc.add_heading('Certifications', level=1)
doc.add_paragraph(
    "- IBM Data Analysis with Python\n"
    "- IBM Data Visualization with Python\n"
    "- IBM SQL and Relational Databases 101\n"
    "- freeCodeCamp Data Analysis with Python"
)

# Projects
doc.add_heading('Projects', level=1)

doc.add_paragraph(
    "Demand vs. Supply Gap Analysis\n"
    "- Analyzed the beverage sales data in India to identify whether demand exceeds supply or if supply meets/exceeds demand.\n"
    "- Created a Power BI dashboard to visualize key metrics and trends, focusing on sales and supply data.\n"
    "- Used calculated fields in Power BI to analyze performance by SKU (Stock Keeping Unit).\n"
    "- Employed Advanced Excel for data analysis and calculations, such as pivot tables, for detailed performance tracking."
)

doc.add_paragraph(
    "Operational Analysis of E-Commerce\n"
    "- Developed a Power BI dashboard for operational analysis of an e-commerce platform, focusing on key areas like product performance, regional sales, shipping metrics, and customer demographics.\n"
    "- Used Advanced Excel to perform data cleaning and analysis, including pivot tables for sales and shipping performance.\n"
    "- Enabled stakeholders to gain actionable insights and improve operational strategies."
)

doc.add_paragraph(
    "AI-Based Security Surveillance System\n"
    "Tools Used: Python, OpenCV, face_recognition, SMTP, PyAutoGUI\n"
    "- Developed a real-time security system that uses facial recognition to identify unauthorized individuals via webcam.\n"
    "- Integrated OpenCV and the face_recognition library to detect faces and match them against known identities.\n"
    "- Implemented automatic screenshot capture, email alerts with image attachments, and audible alarm playback when unknown faces are detected.\n"
    "- Automated surveillance response through Python by combining computer vision, email APIs (SMTP), and user alerts."
)

# Technical Skills
doc.add_heading('Technical Skills', level=1)
doc.add_paragraph(
    "- Programming Languages: Python, SQL\n"
    "- Database Management: Oracle\n"
    "- Data Tools: Power BI, Advanced Excel\n"
    "- Software/Platforms: Visual Studio Code, GitHub"
)

# Languages
doc.add_heading('Languages', level=1)
doc.add_paragraph(
    "- English: Fluent\n"
    "- Malayalam: Native\n"
    "- Tamil: Conversational"
)

# Extracurricular Activities
doc.add_heading('Extracurricular Activities', level=1)
doc.add_paragraph(
    "- College Cricket Team Player: Actively participated in college-level cricket tournaments, demonstrating strong teamwork, discipline, and leadership skills."
)

# Save the document
file_path = "/mnt/data/Sanjay_Resume_Singapore_Malaysia.docx"
doc.save(file_path)

file_path
